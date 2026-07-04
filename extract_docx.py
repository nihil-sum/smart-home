import docx

doc = docx.Document(r'c:\Users\里\Desktop\home\智慧房屋租赁系统 需求规格说明书.docx')

print('=== 段落内容 ===')
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        print(f'[{i}] {p.text}')

print('\n=== 表格内容 ===')
for ti, table in enumerate(doc.tables):
    print(f'\n--- 表格 {ti+1} ---')
    for ri, row in enumerate(table.rows):
        cells = [cell.text.strip() for cell in row.cells]
        print(f'  行{ri}: {" | ".join(cells)}')
